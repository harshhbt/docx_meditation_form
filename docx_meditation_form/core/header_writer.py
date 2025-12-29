from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from docx_meditation_form.core.settings import DocxSettings

APPLICATION_FORM_HEADER = """FORM ‘A’
MEDIATION APPLICATION FORM
[REFER RULE 3(1)]"""

APPLICATION_FORM_SUB_HEADER = """Mumbai District Legal Services Authority
City Civil Court, Mumbai"""


class HeaderWriter:
    """
    Writes the fixed header section of the mediation application form.

    This includes:
    - The main form title (FORM 'A')
    - The sub-header identifying the authority and court

    Parameters
    ----------
    doc : docx.document.Document
        An active python-docx Document instance to write into.
    """

    def __init__(self, doc: Document):
        self.doc = doc

    def write(self):
        """
        Write the complete header to the document.

        This method is idempotent only if called once on a fresh document;
        repeated calls will append duplicate headers.
        """
        self._write_main_header()
        self._write_sub_header()

    def _write_main_header(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(APPLICATION_FORM_HEADER)
        DocxSettings.apply_run_style(run, bold=True)
        p.paragraph_format.space_after = 0

    def _write_sub_header(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(APPLICATION_FORM_SUB_HEADER)
        DocxSettings.apply_run_style(run)
        p.paragraph_format.space_after = Inches(0.34)
