from typing import Iterable, Mapping

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Inches, Length
from docx.table import Table

from docx_meditation_form.core.cell_formatter import CellFormatter
from docx_meditation_form.core.settings import DocxSettings
from docx_meditation_form.dataset.form_values import FormValues


class TableWriter:
    """Writes and formats the main form table."""

    def __init__(self, doc):
        self.doc = doc
        self.settings = DocxSettings(doc)
        self.formatter = CellFormatter(self.settings)

    def init_table(self) -> Table:
        """Create base table structure."""
        return self._create_table(rows=17, columns=3)

    def merge_required_rows(self, table: Table) -> None:
        """Merge cells to match form layout."""
        rules = {
            0: "full",
            2: "right",
            7: "right",
            8: "right",
            14: "full",
            15: "full",
            16: "right",
        }

        for idx, row in enumerate(table.rows):
            mode = rules.get(idx)
            if mode == "full":
                row.cells[0].merge(row.cells[2])
            elif mode == "right":
                row.cells[1].merge(row.cells[2])

    def write_table(self, table: Table, v: FormValues) -> None:
        """
        Populate table row-by-row.

        Each index maps to a logical form section:
        - applicant block
        - defendant block
        - dispute block
        """
        for idx, row in enumerate(table.rows):
            if idx == 0:
                # DETAILS OF PARTIES header
                self.formatter.left_label_cell(row.cells[0], "DETAILS OF PARTIES:")

            elif idx == 1:
                # Applicant name row
                self._write_name_row(
                    row,
                    number="1",
                    label="Name of\nApplicant",
                    value=v.APPLICANT_NAME,
                )

            elif idx == 2:
                # Applicant address subheading
                self.formatter.left_label_cell(
                    row.cells[1], "Address and contact details of Applicant"
                )

            elif 3 <= idx <= 6:
                # Applicant contact block (address, phone, mobile, email)
                self._write_contact_block(row, idx, v, "APPLICANT")

            elif idx == 7:
                # Opposite party section header
                self.formatter.center_label_cell(row.cells[0], "2", bold=False)
                self.formatter.left_label_cell(
                    row.cells[1],
                    "Name, Address and Contact details of Opposite Party:",
                    bold=True,
                )

            elif idx == 8:
                # Defendant address subheading
                self.formatter.left_label_cell(
                    row.cells[1],
                    "Address and contact details of Defendant/s",
                    bold=True,
                )

            elif idx == 9:
                # Defendant name row
                self._write_name_row(
                    row,
                    number=None,
                    label="Name",
                    value=v.DEFENDANT_NAME,
                )

            elif 10 <= idx <= 13:
                # Defendant contact block
                self._write_contact_block(row, idx, v, "DEFENDANT")

            elif idx == 14:
                # DETAILS OF DISPUTE header
                self.formatter.left_label_cell(row.cells[0], "DETAILS OF DISPUTE:")

            elif idx == 15:
                # Dispute title line
                self.formatter.center_label_cell(
                    row.cells[0],
                    "THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES,2018",
                    bold=True,
                    underline=True,
                )

            elif idx == 16:
                # Nature of dispute description
                self.formatter.left_label_cell(
                    row.cells[1],
                    "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):",
                    bold=True,
                )

    def _write_name_row(self, row, *, number, label, value) -> None:
        """Write numbered name row."""
        if number is not None:
            self.formatter.center_label_cell(row.cells[0], number, bold=False)
        self.formatter.left_label_cell(row.cells[1], label, bold=True)
        self.formatter.left_label_cell(row.cells[2], value, bold=True)

    def _write_contact_block(self, row, idx: int, v: FormValues, prefix: str) -> None:
        """
        Write address / phone / email rows.

        idx controls which field is written.
        prefix selects `APPLICANT_*` or `DEFENDANT_*` values.
        """

        def field(name: str) -> str:
            return getattr(v, f"{prefix}_{name}")

        if idx in (3, 10):
            # Address block
            if idx == 3:
                self.formatter.center_label_cell(row.cells[0], "1", bold=False)

            self.formatter.left_label_cell(row.cells[1], "Address", bold=True)

            cell = row.cells[2]

            p0 = cell.paragraphs[0]
            self.formatter.apply_paragraph(
                p0,
                "REGISTERED ADDRESS:\n",
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            self.formatter.apply_paragraph(
                p0,
                field("BRANCH_ADDRESS"),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Inches(0.28 if idx == 3 else 0.68),
            )

            p1 = cell.add_paragraph()
            self.formatter.apply_paragraph(
                p1,
                "CORRESPONDENCE BRANCH ADDRESS:\n",
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            self.formatter.apply_paragraph(
                p1,
                field("CORRESPONDENCE_BRANCH_ADDRESS"),
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )

        elif idx in (4, 11):
            # Telephone
            self._simple_value_row(row, "Telephone No.", field("PHONE"))

        elif idx in (5, 12):
            # Mobile
            self._simple_value_row(row, "Mobile No.", field("MOBILE"))

        elif idx in (6, 13):
            # Email
            self._simple_value_row(row, "Email ID", field("EMAIL_ID"))

    def _simple_value_row(self, row, label: str, value: str) -> None:
        """Label + value row."""
        self.formatter.left_label_cell(row.cells[1], label, bold=True)
        self.formatter.left_label_cell(row.cells[2], value, bold=True)

    def set_columns_width(self, table: Table, widths: list[Inches]) -> None:
        """Apply column widths."""
        for idx, width in enumerate(widths):
            table.columns[idx].width = width
            for row in table.rows:
                row.cells[idx].width = width

    def set_rows_height(
        self, table: Table, rows_height_dataset: Mapping[str, Length]
    ) -> None:
        """Apply row heights using index-based rules."""
        rules = [
            (lambda i: i == 2 or 4 <= i <= 8, "small_text"),
            (lambda i: i == 3, "section"),
            (lambda i: i == 10, "major_section"),
            (lambda i: i == 16, "tiny_gap"),
        ]

        for idx, row in enumerate(table.rows):
            key = self._resolve_height_key(idx, rules, default="base_text")
            self.set_row_height(row, rows_height_dataset[key])

    def _resolve_height_key(
        self,
        idx: int,
        rules: Iterable,
        *,
        default: str,
    ) -> str:
        """Resolve height key for a row index."""
        for predicate, key in rules:
            if predicate(idx):
                return key
        return default

    def _create_table(self, rows: int, columns: int) -> Table:
        """Create table with borders."""
        table = self.doc.add_table(rows, columns)
        self.set_table_borders(table, "000000")
        return table

    def set_table_borders(self, table: Table, color: str) -> None:
        """Apply uniform table borders."""
        tbl = table._tbl
        borders = OxmlElement("w:tblBorders")

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(ns.qn("w:val"), "single")
            e.set(ns.qn("w:sz"), "4")
            e.set(ns.qn("w:space"), "0")
            e.set(ns.qn("w:color"), color)
            borders.append(e)

        tbl.tblPr.append(borders)

    def set_row_height(self, row, height: Length) -> None:
        """
        Set minimum row height.

        :param row: `docx.table.Table._Row`
        :param height: `docx.shared.Length`
        
        Thanks: https://stackoverflow.com/a/43407902/32112205
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()

        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), str(height.twips))
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)
