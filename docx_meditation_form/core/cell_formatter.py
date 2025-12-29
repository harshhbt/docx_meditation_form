from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length
from docx.text.paragraph import Paragraph

from docx_meditation_form.core.settings import DocxSettings


class CellFormatter:
    """Utility for formatting table cells and paragraphs."""

    def __init__(self, settings: DocxSettings):
        """Create formatter with shared document settings."""
        self.settings = settings

    def center_label_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        underline: bool = False,
    ) -> None:
        """
        Write centered label text in a table cell.

        :param cell: docx.table._Cell
        :param text: Cell text
        """
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        self.apply_paragraph(
            p, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold, underline=underline
        )

    def top_label_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        underline: bool = False,
    ) -> None:
        """
        Write top-aligned label text in a table cell.

        :param cell: docx.table._Cell
        :param text: Cell text
        """
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        self.apply_paragraph(
            p, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=bold, underline=underline
        )

    def left_label_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = True,
    ) -> None:
        """
        Write left-aligned label text in a table cell.

        :param cell: docx.table._Cell
        :param text: Cell text
        """
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        self.apply_paragraph(
            p,
            text,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            bold=bold,
        )

    def apply_paragraph(
        self,
        p: Paragraph,
        text: str,
        *,
        align: WD_ALIGN_PARAGRAPH | None = None,
        bold: bool = False,
        underline: bool = False,
        space_before: Length | int = 0,
        space_after: Length | int = 0,
    ) -> None:
        """
        Apply text and formatting to a paragraph.

        :param p: docx.text.paragraph.Paragraph
        :param text: Paragraph text
        """
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after

        run = p.add_run(text)
        self.settings.apply_run_style(
            run,
            bold=bold,
            underline=underline,
        )
